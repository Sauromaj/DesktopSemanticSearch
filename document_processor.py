"""
Document Processor Module
Handles extraction of text from various document formats (PDF, Word, Excel)
and processing into embeddings for semantic search.
"""
import os
import time
from pathlib import Path
from typing import List, Dict, Any, Generator, Tuple
import logging
import hashlib

# Document parsing libraries
import pypdf
import docx
import pandas as pd
from openpyxl import load_workbook

# Progress tracking
from rich.progress import Progress, TextColumn, BarColumn, TimeElapsedColumn, TimeRemainingColumn

# Local modules
from config import ApplicationConfig
from ui_manager import UIManager

logger = logging.getLogger("semantic_search")

class DocumentProcessor:
    """Class to handle document processing and indexing"""
    
    def __init__(self, config: ApplicationConfig, ui_manager: UIManager):
        """Initialize the document processor with configuration"""
        self.config = config
        self.ui_manager = ui_manager
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Note: .doc needs additional handling in implementation
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv
        }
        
    def index_directory(self, directory_path: Path, force: bool = False) -> bool:
        """
        Index all documents in the specified directory
        
        Args:
            directory_path: Path to the directory to index
            force: Whether to force reindexing of already indexed documents
            
        Returns:
            bool: True if indexing was successful
        """
        # Get list of files to index
        files_to_index = self._get_indexable_files(directory_path)
        
        if not files_to_index:
            self.ui_manager.display_warning(f"No supported documents found in {directory_path}")
            return False
            
        # Check if we need to process them
        if not force:
            # Filter out already indexed files based on modification time and hash
            index_metadata = self._load_index_metadata()
            files_to_index = [f for f in files_to_index if self._needs_indexing(f, index_metadata)]
            
        if not files_to_index:
            self.ui_manager.display_message("All documents are already indexed and up to date")
            return True
            
        # Index the files
        total_files = len(files_to_index)
        self.ui_manager.display_message(f"Indexing {total_files} documents...")
        
        processed_documents = []
        
        # Setup progress bar
        with Progress(
            TextColumn("[bold blue]{task.description}"),
            BarColumn(bar_width=40),
            "[progress.percentage]{task.percentage:>3.0f}%",
            TimeElapsedColumn(),
            TimeRemainingColumn(),
            console=self.ui_manager.console
        ) as progress:
            
            task = progress.add_task(f"[cyan]Processing documents...", total=total_files)
            
            for file_path in files_to_index:
                try:
                    # Extract text and metadata from document
                    doc_content = self._process_document(file_path)
                    if doc_content:
                        processed_documents.append(doc_content)
                        
                        # Update metadata for this file
                        file_hash = self._get_file_hash(file_path)
                        self._update_index_metadata(file_path, file_hash)
                    
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
                    self.ui_manager.display_warning(f"Skipped {file_path}: {str(e)}")
                
                # Update progress
                progress.update(task, advance=1)
        
        # Store processed documents in the vector database
        if processed_documents:
            self.ui_manager.display_message(f"Creating embeddings for {len(processed_documents)} documents...")
            from vector_store import VectorStore
            vector_store = VectorStore(self.config.vector_db_path)
            vector_store.add_documents(processed_documents)
            
        return True
    
    def _get_indexable_files(self, directory_path: Path) -> List[Path]:
        """Get all files that can be indexed from the directory"""
        indexable_files = []
        
        for file_path in directory_path.glob('**/*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                indexable_files.append(file_path)
                
        return indexable_files
    
    def _needs_indexing(self, file_path: Path, index_metadata: Dict[str, Dict[str, Any]]) -> bool:
        """Check if a file needs to be indexed based on modification time and hash"""
        file_str = str(file_path)
        
        # If file is not in metadata, it needs indexing
        if file_str not in index_metadata:
            return True
            
        # Check if file has been modified since last indexing
        last_modified = file_path.stat().st_mtime
        if last_modified > index_metadata[file_str].get('last_indexed', 0):
            return True
            
        # Check if content has changed using hash
        current_hash = self._get_file_hash(file_path)
        if current_hash != index_metadata[file_str].get('hash', ''):
            return True
            
        return False
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Calculate a hash of file contents for change detection"""
        hash_md5 = hashlib.md5()
        
        with open(file_path, "rb") as f:
            # Read in chunks to avoid memory issues with large files
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
                
        return hash_md5.hexdigest()
    
    def _load_index_metadata(self) -> Dict[str, Dict[str, Any]]:
        """Load metadata about indexed files"""
        import json
        
        metadata_path = Path(self.config.data_dir) / "index_metadata.json"
        
        if not metadata_path.exists():
            return {}
            
        try:
            with open(metadata_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Error loading index metadata: {str(e)}")
            return {}
    
    def _update_index_metadata(self, file_path: Path, file_hash: str) -> None:
        """Update metadata for an indexed file"""
        import json
        
        metadata_path = Path(self.config.data_dir) / "index_metadata.json"
        metadata = self._load_index_metadata()
        
        # Update metadata for this file
        metadata[str(file_path)] = {
            'last_indexed': time.time(),
            'hash': file_hash,
            'size': file_path.stat().st_size
        }
        
        # Ensure data directory exists
        os.makedirs(Path(self.config.data_dir), exist_ok=True)
        
        # Save updated metadata
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)
    
    def _process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process a document and extract text and metadata"""
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
            
        processor_func = self.supported_extensions[extension]
        text_content = processor_func(file_path)
        
        if not text_content:
            logger.warning(f"No content extracted from {file_path}")
            return None
            
        # Create document record
        document = {
            'path': str(file_path),
            'filename': file_path.name,
            'extension': extension,
            'content': text_content,
            'last_modified': file_path.stat().st_mtime,
            'size': file_path.stat().st_size
        }
        
        return document
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file"""
        text_content = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
                    
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise ValueError(f"PDF parsing error: {str(e)}")
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from a Word document (.docx)"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                text_content.append(para.text)
                
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content.append(" | ".join(row_text))
                    
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading Word document {file_path}: {str(e)}")
            raise ValueError(f"Word document parsing error: {str(e)}")
    
    def _process_excel(self, file_path: Path) -> str:
        """Extract text from an Excel file"""
        try:
            workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                # Process each row
                for row in sheet.iter_rows():
                    row_values = [str(cell.value) if cell.value is not None else "" for cell in row]
                    text_content.append(" | ".join(row_values))
                    
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            raise ValueError(f"Excel parsing error: {str(e)}")
    
    def _process_csv(self, file_path: Path) -> str:
        """Extract text from a CSV file"""
        try:
            df = pd.read_csv(file_path)
            text_content = []
            
            # Add header
            headers = " | ".join(df.columns.astype(str))
            text_content.append(headers)
            
            # Add rows
            for _, row in df.iterrows():
                row_values = " | ".join(row.astype(str))
                text_content.append(row_values)
                
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {str(e)}")
            raise ValueError(f"CSV parsing error: {str(e)}")
